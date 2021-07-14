# import the necessary packages
from imutils.video import FPS
import numpy as np
import argparse
import imutils
import time
import cv2
import os
import docx
from docx import Document
from docx.shared import Cm
from tkinter.filedialog import askopenfilename
from PIL import Image
# construct the argument parse and parse the arguments
ap = argparse.ArgumentParser()
document = Document()
p = document.add_paragraph()
r = p.add_run()
detected_ids = set()
"""ap.add_argument("-i", "--input", required=True,
	help="path to input video")"""
ap.add_argument("-c", "--confidence", type=float, default=0.5,
	help="minimum probability to filter weak detections")
ap.add_argument("-t", "--threshold", type=float, default=0.3,
	help="threshold when applyong non-maxima suppression")
# ap.add_argument("-v", "--video", required=True,
# 	help="path to input video file")
args = vars(ap.parse_args())

def save_picture(picture, path):
	image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
	picture = Image.fromarray(image)
	picture.save(path)

def add_picture_to_doc(r, picture, width_cm=15, height_cm=10):
    path = 'last.jpg'
    save_picture(picture, path)
    r.add_picture(path, width=Cm(width_cm), height=Cm(height_cm))


def add_to_doc(r, frame, text, num_without_ladder, num_person, num_with_ladder):
    add_picture_to_doc(r, frame)
    r.add_text(text)
    r.add_break()

def save_report(filename):
	print('Сохраняем отчет...')
	if not os.path.isdir('Отчеты'):
		os.mkdir('Отчеты')    
	if filename:
		doc_filename = filename.replace('.mp4', '.docx').replace('Видео/','Отчеты/')
		document.save('' + doc_filename)
		print('Отчет сохранен в ' + doc_filename)

# load the COCO class labels our YOLO model was trained on
labelsPath = "classes.txt"
LABELS = open(labelsPath).read().strip().split("\n")

# initialize a list of colors to represent each possible class label
# np.random.seed(42)
np.random.seed(100)
COLORS = np.random.randint(0, 255, size=(len(LABELS), 3),dtype="uint8")

# derive the paths to the YOLO weights and model configuration
weightsPath = "yolov3_training_last.weights"
configPath = "yolov3_testing.cfg"

# load our YOLO object detector trained on COCO dataset (80 classes)
# and determine only the *output* layer names that we need from YOLO
print("[INFO] loading YOLO from disk...")
net = cv2.dnn.readNetFromDarknet(configPath, weightsPath)
ln = net.getLayerNames()
ln = [ln[i[0] - 1] for i in net.getUnconnectedOutLayers()]

# initialize the video stream, pointer to output video file, and
# frame dimensions
filename = askopenfilename() # show an "Open" dialog box and return the path to the selected file
vs = cv2.VideoCapture(filename)
# fps = FPS().start()
(W, H) = (None, None)

# loop over frames from the video file stream
while True:
	# read the next frame from the file
	(grabbed, frame) = vs.read()
	frame = cv2.resize(frame,(1280,720))
	# frame = imutils.resize(frame, width=450)
	ff = frame
	frame = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
	frame = np.dstack([frame, frame, frame])

	# if the frame was not grabbed, then we have reached the end
	# of the stream
	if not grabbed:
		break

	# if the frame dimensions are empty, grab them
	if W is None or H is None:
		(H, W) = frame.shape[:2]

	# construct a blob from the input frame and then perform a forward
	# pass of the YOLO object detector, giving us our bounding boxes
	# and associated probabilities
	blob = cv2.dnn.blobFromImage(frame, 1 / 255.0, (416, 416),
		swapRB=True, crop=False)
	net.setInput(blob)
	start = time.time()
	layerOutputs = net.forward(ln)
	end = time.time()

	# initialize our lists of detected bounding boxes, confidences,
	# and class IDs, respectively
	boxes = []
	confidences = []
	classIDs = []

	# loop over each of the layer outputs
	for output in layerOutputs:
		# loop over each of the detections
		for detection in output:
			# extract the class ID and confidence (i.e., probability)
			# of the current object detection
			scores = detection[5:]
			classID = np.argmax(scores)
			confidence = scores[classID]

			# filter out weak predictions by ensuring the detected
			# probability is greater than the minimum probability
			if confidence > args["confidence"]:
				# scale the bounding box coordinates back relative to
				# the size of the image, keeping in mind that YOLO
				# actually returns the center (x, y)-coordinates of
				# the bounding box followed by the boxes' width and
				# height
				box = detection[0:4] * np.array([W, H, W, H])
				(centerX, centerY, width, height) = box.astype("int")

				# use the center (x, y)-coordinates to derive the top
				# and and left corner of the bounding box
				x = int(centerX - (width / 2))
				y = int(centerY - (height / 2))

				# update our list of bounding box coordinates,
				# confidences, and class IDs
				boxes.append([x, y, int(width), int(height)])
				confidences.append(float(confidence))
				classIDs.append(classID)
			
	# apply non-maxima suppression to suppress weak, overlapping
	# bounding boxes
	idxs = cv2.dnn.NMSBoxes(boxes, confidences, args["confidence"],
		args["threshold"])
	num_with_ladder = 0
	num_without_ladder = 0
	num_person = 0
	

	# ensure at least one detection exists
	if len(idxs) > 0:
		# loop over the indexes we are keeping
		for i in idxs.flatten():
			# extract the bounding box coordinates
			(x, y) = (boxes[i][0], boxes[i][1])
			(w, h) = (boxes[i][2], boxes[i][3])
							
			# draw a bounding box rectangle and label on the frame
			color = [int(c) for c in COLORS[classIDs[i]]]
			cv2.rectangle(ff, (x, y), (x + w, y + h), color, 2)
			text = "{}: {:.4f}".format(LABELS[classIDs[i]],
				confidences[i])

	

			cv2.putText(ff, text, (x, y - 5),
				cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

			if classIDs[i] == 0:
				num_with_ladder += 1


			if classIDs[i] == 1:
				num_person += 1

			if classIDs[i] == 2:
				num_without_ladder += 1

		# save when a violation is detected
		if num_without_ladder > 0:
			text = '\n  Вид нарушения:  Спуск-подъем персонала на площадку (кузов) транспортного средства осуществляется без использования лестницы (стремянок, передвижных площадок и т.п.) или специальной площадки. '
			add_to_doc(r, ff, text, num_without_ladder, num_person, num_with_ladder)
			save_report(filename)
			# break
	cv2.rectangle(ff, (0, 900), (340, 610), (0,0,0), -1)
	textcadr1 = f"Amount of people: {num_with_ladder}"
	textcadr2 = f"Person without ladder: {num_without_ladder}"
	textcadr3 = f"Person with ladder: {num_person}"
	cv2.putText(ff, textcadr1, (10,640), cv2.FONT_HERSHEY_SIMPLEX, 0.80, (0,255,0), 2)
	cv2.putText(ff, textcadr2, (10,700), cv2.FONT_HERSHEY_SIMPLEX, 0.80, (0,0,255), 2)
	cv2.putText(ff, textcadr3, (10,670), cv2.FONT_HERSHEY_SIMPLEX, 0.80, (255,0,0), 2)
	cv2.imshow("output",ff)

	key = cv2.waitKey(1) & 0xFF
	# if the `s` key was pressed, save the report
	if key == ord("s"):
		if num_without_ladder >= 1:
			text = '\n  Вид нарушения:  Спуск-подъем персонала на площадку (кузов) транспортного средства осуществляется без использования лестницы (стремянок, передвижных площадок и т.п.) или специальной площадки. '
			add_to_doc(r, ff, text, num_without_ladder, num_person, num_with_ladder)
			save_report(filename)
		else:
			text = '\n Нарушений не выявлено'
			add_to_doc(r, ff, text, num_with_ladder, num_person, num_without_ladder)
			save_report(filename)
			if key == ord("w"):
				print("something")

	# if the `q` key was pressed, break from the loop
	if key == ord("q"):
		break
# 			fps.update()
# 			fps.stop()
# print("[INFO] elasped time: {:.2f}".format(fps.elapsed()))
# print("[INFO] approx. FPS: {:.2f}".format(fps.fps()))
# release the file pointers
print("[INFO] cleaning up...")
vs.release()
